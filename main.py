from docx.shared import RGBColor
import string
import docx

doc = docx.Document('./test.docx')
f = open('./text.txt', 'r')
to_hide_text = "".join(f.readlines())
color = RGBColor(8, 92, 168)


def get_letters_to_hide(t):
    return list("".join(t.translate(str.maketrans('', '', string.punctuation)).split()))


def hide_text(document, text_to_hide, filename):
    letters_to_hide = get_letters_to_hide(text_to_hide)
    amount_of_letters = len(letters_to_hide)
    for para in document.paragraphs:
        text = para.text.split()
        para.text = ''
        for word in text:
            if len(letters_to_hide) == 0:
                para.add_run(word + " ")
                continue
            if letters_to_hide[0].lower() in word.lower():
                words_letters = list(word)
                # print(words_letters)
                # is_marked = False
                for l in words_letters:
                    if len(letters_to_hide) != 0 and l.lower() == letters_to_hide[0].lower():
                        isLower = l.islower()
                        if not isLower:
                            l = l.upper()
                        del letters_to_hide[0]
                        # print(letters_to_hide)
                        para.add_run(l).font.color.rgb = color

                    else:
                        para.add_run(l)
                para.add_run(" ")
            else:
                para.add_run(word + " ")
    document.save(filename)
    print(f"{amount_of_letters - len(letters_to_hide)} of {amount_of_letters} letters have been hidden!")


hide_text(doc, to_hide_text, 'out.docx')